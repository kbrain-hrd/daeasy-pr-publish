"""테스트용 홍보자료 양식을 만든다.

AI 챔피언 그린 5회차 건. 담당자가 실제로 작성했다고 가정한 내용이며,
수치는 운영 EMS(kbrain-ems.vercel.app) 실제 값으로 맞췄다.
"""
import os
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from prpub.parse import docx_value_cells

ROOT = Path(__file__).resolve().parent.parent
FOLDER = ROOT / "test" / "접수함" / "2026-08-25_AI챔피언그린_5회차"

# 비대면 과정이라 강의실 사진은 쓰지 않는다. 담당자가 산출물 화면을 냈다고 본다.
PHOTOS = [
    ROOT / "docs/preview/img/vulnerable-area-dashboard.png",
    ROOT / "docs/preview/img/civil-complaint-ai.png",
]

DATA = {
    "course_name": "AI 챔피언 그린 5회차",
    "course": "AI 챔피언 그린",
    "org": "공모 선발 (공공기관)",
    "dates": "2026-08-04 ~ 2026-08-25",
    "round": "5회차",
    "participants": "107명 선발 (정원 100명 · 신청 615명 · 경쟁률 6.2:1), 공공기관 주무관 등 비개발 실무자",
    "topics": (
        "데이터 읽기 — 부서에서 쓰는 엑셀 파일을 직접 열어 빠진 값과 중복을 찾고, "
        "그 숫자가 무엇을 뜻하는지 따져봤습니다.\n"
        "생성형 AI로 업무 문서 다루기 — 보고서 초안과 회의록을 AI로 정리하고, "
        "그대로 쓰면 안 되는 부분을 골라내는 연습을 했습니다.\n"
        "노코드로 프로토타입 만들기 — 아이디어를 문서로 남기지 않고 실제로 클릭되는 화면으로 만들었습니다. "
        "코드는 한 줄도 쓰지 않습니다.\n"
        "수행평가 — 각자 부서의 문제를 하나 골라 기획안과 프로토타입으로 제출했습니다."
    ),
    "process": (
        "선수과목(기초역량) → 사전 온라인 3일 → OT와 집중수업(비대면, 과제 중심 실습) → "
        "셀프스터디 2일 → 수행평가 → 인증(75점 이상). 총 41시간. 전 과정 비대면."
    ),
    "outputs": (
        "참여자별 기획안과 업무 프로토타입. 민원 분류 도구, 지역 데이터 대시보드 등 "
        "부서 업무에 바로 붙일 수 있는 형태로 나왔습니다."
    ),
    "highlight": (
        "샘플 데이터를 쓰지 않습니다. 참여자가 실제로 업무에 쓰는 부서 데이터를 그대로 가져와 실습합니다.\n"
        "기획안에서 멈추지 않고 클릭되는 프로토타입까지 만듭니다. 교육이 끝나면 결과물이 남습니다.\n"
        "비대면인데 평균 출석률이 100%였습니다. OT부터 마지막 수업까지 빈자리가 없었습니다.\n"
        "선발 107명 중 96명이 수료하고 92명이 인증을 받았습니다."
    ),
    "satisfaction": "8.6 / 10점 (88명 응답, 응답률 82%)",
    "headline": "코드 한 줄 없이 만든 부서 업무 도구",
    "hashtags": "#AI챔피언 #공공데이터 #노코드 #공무원교육 #행정안전부",
}

shutil.rmtree(FOLDER, ignore_errors=True)
os.makedirs(FOLDER / "사진", exist_ok=True)

doc = Document(str(ROOT / "templates/홍보자료_양식.docx"))
n = 0
for f, cell in docx_value_cells(doc):
    if f.key not in DATA:
        continue
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    lines = DATA[f.key].split("\n")
    cell.paragraphs[0].text = lines[0]
    for extra in lines[1:]:
        cell.add_paragraph(extra)
    n += 1
doc.save(str(FOLDER / "홍보자료_양식.docx"))

for i, src in enumerate(PHOTOS, 1):
    shutil.copy2(src, FOLDER / "사진" / f"{i:02d}{src.suffix}")

print(f"{n}개 항목 채움 · 사진 {len(PHOTOS)}장 -> {FOLDER.relative_to(ROOT)}")
