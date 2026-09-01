"""양식 파일 생성: docx 를 만들고, 한글이 설치돼 있으면 hwp 로도 변환한다.

디자인은 daeasy.vercel.app 톤 — 흰 바탕, 얇은 회색 구분선, 검정 제목, 포인트 블루(#2563EB).
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .schema import FIELDS, SECTIONS

FONT = "맑은 고딕"
INK = RGBColor(0x0F, 0x0F, 0x0F)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
ZINC_500 = RGBColor(0x71, 0x71, 0x7A)
ZINC_400 = RGBColor(0x9F, 0x9F, 0xA9)
LINE = "E4E4E7"  # zinc-200
LINE_DARK = "0F0F0F"
FILL = "FAFAFA"  # zinc-50, 입력칸 바탕

HALF_W = Cm(8.3)
GAP_W = Cm(0.4)
VALUE_H = Cm(0.85)  # 입력칸 최소 높이
VALUE_H_MULTI = Cm(2.2)


def _row_height(cell, height) -> None:
    """셀이 속한 행의 최소 높이."""
    tr = cell._tc.getparent()
    tr_pr = tr.get_or_add_trPr()
    el = OxmlElement("w:trHeight")
    el.set(qn("w:val"), str(int(height.twips)))
    el.set(qn("w:hRule"), "atLeast")
    tr_pr.append(el)


def _run(paragraph, text: str, size: float, bold: bool = False, color: RGBColor = INK, spacing: int = 0):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    if spacing:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        r._element.rPr.append(sp)
    return r


def _para_border_bottom(paragraph, color: str, sz: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color)
    bdr.append(b)
    p_pr.append(bdr)


def _cell_borders(cell, bottom: tuple[str, int] | None) -> None:
    """좌·우·상 선 없음, 아래 선만 (색, 굵기) 로."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    el = OxmlElement("w:bottom")
    if bottom:
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(bottom[1]))
        el.set(qn("w:color"), bottom[0])
    else:
        el.set(qn("w:val"), "nil")
    borders.append(el)
    tc_pr.append(borders)


def _shade(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _table_frame(table, indent: int) -> None:
    """표 자체의 테두리 없음(한글 변환 시 기본 외곽선이 생기는 것 방지) + 왼쪽 들여쓰기 보정."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(indent))
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)


def _cell_margins(table, top: int, bottom: int, left: int, right: int) -> None:
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def _no_split(row) -> None:
    """행이 페이지 사이에서 잘리지 않게."""
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    tr_pr.append(el)


def _tight(paragraph, before: float = 0, after: float = 0, line: float = 1.15) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def build_docx(out: Path) -> Path:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.6)

    # ── 머리 ──
    p = doc.add_paragraph()
    _tight(p, after=2)
    _run(p, "DAEASY", 8.5, bold=True, color=ACCENT, spacing=60)

    p = doc.add_paragraph()
    _tight(p, after=6)
    _run(p, "홍보자료 등록 양식", 20, bold=True)
    _para_border_bottom(p, LINE_DARK, 8)

    # ── 표: 항목명 행 + 입력 행이 한 쌍. 세 칸 [A | 여백 | B] 로 짧은 항목은 좌우 2단 ──
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _cell_margins(table, top=50, bottom=50, left=80, right=80)
    _table_frame(table, indent=-80)
    for col, w in zip(table.columns, (HALF_W, GAP_W, HALF_W)):
        col.width = w

    def section_into(cell, key: str) -> None:
        """구분 제목을 첫 항목 칸 안에 넣어 항목과 항상 같은 쪽에 있게 한다 (한글은 keepNext 를 무시)."""
        title, desc = SECTIONS[key]
        num, _, name = title.partition(" ")
        mp = cell.paragraphs[0]
        _tight(mp, before=12, after=6)
        _run(mp, num, 9, bold=True, color=ACCENT, spacing=20)
        _run(mp, "   " + name.strip(), 10.5, bold=True)
        if desc:
            _run(mp, "    " + desc, 8, color=ZINC_500)
        _para_border_bottom(mp, LINE_DARK, 6)

    def label_into(cell, f, section_key: str | None = None) -> None:
        cell.text = ""
        if section_key:
            section_into(cell, section_key)
            lp = cell.add_paragraph()
        else:
            lp = cell.paragraphs[0]
        _tight(lp, before=5)
        _run(lp, f.label, 9, bold=True)
        guide = " / ".join(f.choices) if f.choices else f.hint
        if guide:
            gp = cell.add_paragraph()
            _tight(gp, before=1)
            _run(gp, guide, 7, color=ZINC_400)
        _cell_borders(cell, None)

    def value_into(cell, f) -> None:
        cell.text = ""
        _tight(cell.paragraphs[0])
        _shade(cell, FILL)
        _cell_borders(cell, (LINE, 4))
        _row_height(cell, VALUE_H_MULTI if f.multiline else VALUE_H)

    def field_pair(fields: list) -> None:
        """fields 가 1개면 전체 폭, 2개면 좌우 2단. 구분의 첫 항목이면 제목을 그 칸 위에 넣는다."""
        section_key = fields[0].key if fields[0].key in SECTIONS else None
        lrow = table.add_row()
        vrow = table.add_row()
        _no_split(lrow)
        _no_split(vrow)
        for row in (lrow, vrow):
            row.cells[0].width = HALF_W
            row.cells[1].width = GAP_W
            row.cells[2].width = HALF_W
        if len(fields) == 1:
            lc = lrow.cells[0].merge(lrow.cells[2])
            vc = vrow.cells[0].merge(vrow.cells[2])
            label_into(lc, fields[0], section_key)
            value_into(vc, fields[0])
            return
        for col, f in zip((0, 2), fields):
            label_into(lrow.cells[col], f)
            value_into(vrow.cells[col], f)
        for row in (lrow, vrow):
            row.cells[1].text = ""
            _cell_borders(row.cells[1], None)

    queue: list = []
    for f in FIELDS:
        if f.key in SECTIONS and queue:
            field_pair(queue)
            queue = []
        # 구분의 첫 항목은 제목을 품으므로 항상 전체 폭
        if f.half and f.key not in SECTIONS:
            queue.append(f)
            if len(queue) == 2:
                field_pair(queue)
                queue = []
        else:
            if queue:
                field_pair(queue)
                queue = []
            field_pair([f])
    if queue:
        field_pair(queue)

    _tight(doc.add_paragraph())  # 표 뒤 빈 문단 (표로 문서가 끝나지 않게)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def convert_with_hwp(src: Path, fmt: str, out: Path) -> Path:
    """한글 COM 자동화로 변환. fmt: 'HWP' | 'HWPX' | 'PDF'. 한글이 없으면 예외."""
    import win32com.client  # noqa: PLC0415

    hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
    try:
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        hwp.XHwpWindows.Item(0).Visible = False
        if not hwp.Open(str(src.resolve()), "", "forceopen:true"):
            raise RuntimeError(f"한글에서 열지 못함: {src}")
        hwp.SaveAs(str(out.resolve()), fmt, "")
    finally:
        hwp.Quit()
    return out
