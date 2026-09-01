"""양식 파일(docx / hwpx / hwp) → {key: value} 딕셔너리.

양식 구조: 표의 한 행에 항목명(+안내)이 있고, 바로 아래 행이 그 입력칸이다.
짧은 항목은 [A | 여백 | B] 세 칸으로 좌우 2단. 항목명 행과 입력 행은 칸 구조가 같으므로
같은 칸 위치(index)로 짝을 맞춘다.
"""

import re
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from .schema import BY_LABEL, Field, normalize_label


def _clean_value(raw: str, choices: tuple[str, ...]) -> str:
    v = raw.replace("\r", "")
    if choices:
        # "교육후기 " 나 "→ 교육후기" 처럼 선택지 하나만 남긴 경우 그 선택지로 정규화
        hits = [c for c in choices if c in v]
        if len(hits) == 1 and v.replace(hits[0], "").strip(" /→:·\n") == "":
            return hits[0]
    lines = [ln.rstrip() for ln in v.split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines).strip()


def pair_rows(rows: list[list[str]]) -> list[tuple[Field, int, int]]:
    """행별 칸 텍스트 목록 → (Field, 입력행 index, 칸 index). 입력칸 위치를 알려준다."""
    out: list[tuple[Field, int, int]] = []
    pending: list[tuple[int, Field]] = []
    for r, cells in enumerate(rows):
        labels = [(i, BY_LABEL.get(normalize_label(t))) for i, t in enumerate(cells) if t.strip()]
        if labels and all(f is not None for _, f in labels):
            pending = [(i, f) for i, f in labels if f is not None]
            continue
        if pending:
            for i, f in pending:
                if i < len(cells):
                    out.append((f, r, i))
            pending = []
    return out


def _rows_to_data(rows: list[list[str]]) -> dict[str, str]:
    return {f.key: _clean_value(rows[r][i], f.choices) for f, r, i in pair_rows(rows)}


# ── docx ──


def docx_rows(table) -> list[list]:
    """행별 셀 객체 목록 (병합 셀은 한 번만)."""
    result = []
    for row in table.rows:
        seen: set[int] = set()
        cells = []
        for c in row.cells:
            if id(c._tc) in seen:
                continue
            seen.add(id(c._tc))
            cells.append(c)
        result.append(cells)
    return result


def docx_value_cells(doc) -> list[tuple[Field, object]]:
    """(Field, 입력 셀) 목록 — 파싱과 테스트 픽스처 채우기가 공유."""
    out = []
    for table in doc.tables:
        cells_by_row = docx_rows(table)
        texts = [[c.text for c in row] for row in cells_by_row]
        for f, r, i in pair_rows(texts):
            out.append((f, cells_by_row[r][i]))
    return out


def parse_docx(path: Path) -> dict[str, str]:
    doc = Document(str(path))
    data: dict[str, str] = {}
    for table in doc.tables:
        texts = [[c.text for c in row] for row in docx_rows(table)]
        data.update(_rows_to_data(texts))
    return data


# ── hwpx ──

_NS = {"hp": "http://www.hancom.co.kr/hwpml/2011/paragraph"}


def _t_text(t: ET.Element) -> str:
    """<hp:t>앞<hp:lineBreak/>뒤</hp:t> 처럼 줄바꿈 요소 뒤의 tail 텍스트까지 잇는다."""
    parts = [t.text or ""]
    for child in t:
        if child.tag.endswith("lineBreak"):
            parts.append("\n")
        parts.append(child.tail or "")
    return "".join(parts)


def _cell_text(tc: ET.Element) -> str:
    paras = []
    for p in tc.findall(".//hp:p", _NS):
        paras.append("".join(_t_text(t) for t in p.findall(".//hp:t", _NS)))
    return "\n".join(paras)


def parse_hwpx(path: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    with zipfile.ZipFile(path) as z:
        names = sorted(n for n in z.namelist() if re.match(r"Contents/section\d+\.xml$", n))
        for name in names:
            root = ET.fromstring(z.read(name))
            for tbl in root.iter(f"{{{_NS['hp']}}}tbl"):
                texts = [[_cell_text(tc) for tc in tr.findall("hp:tc", _NS)] for tr in tbl.findall("hp:tr", _NS)]
                data.update(_rows_to_data(texts))
    return data


def parse_hwp(path: Path) -> dict[str, str]:
    """구형 .hwp 는 한글 COM 으로 hwpx 로 바꾼 뒤 파싱한다 (한글 설치 필요)."""
    from .template import convert_with_hwp  # noqa: PLC0415

    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td) / (path.stem + ".hwpx")
        convert_with_hwp(path, "HWPX", tmp)
        return parse_hwpx(tmp)


def parse_form(path: Path) -> dict[str, str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".hwpx":
        return parse_hwpx(path)
    if ext == ".hwp":
        return parse_hwp(path)
    raise ValueError(f"지원하지 않는 양식 형식: {path.name}")
